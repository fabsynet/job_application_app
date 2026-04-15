"""Phase 5 email builder — pure helpers, no I/O except DOCX read."""
from __future__ import annotations

import re
from email.message import EmailMessage
from pathlib import Path

# Non-noreply email regex (order matters: filters first)
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_NOREPLY_PREFIXES = (
    "noreply",
    "no-reply",
    "donotreply",
    "do-not-reply",
    "notifications",
)

_DOCX_MAINTYPE = "application"
_DOCX_SUBTYPE = "vnd.openxmlformats-officedocument.wordprocessingml.document"


def _slug_ascii(value: str) -> str:
    """Strict ASCII slug for filename components: ``[A-Za-z0-9_]`` only.

    Empty / whitespace / pure-unicode inputs fall back to ``"Unknown"`` so
    attachment headers always land in a safe range even when the upstream
    profile or company field is missing or non-ASCII.
    """
    cleaned = (value or "").strip().replace(" ", "_")
    slug = re.sub(r"[^A-Za-z0-9_]", "", cleaned)
    return slug or "Unknown"


def build_subject(*, role: str, company: str) -> str:
    """Locked format: ``Application for {role} at {company}``."""
    role = (role or "Application").strip() or "Application"
    company = (company or "Unknown").strip() or "Unknown"
    return f"Application for {role} at {company}"


def build_attachment_filename(*, full_name: str, company: str) -> str:
    """Locked format: ``{FullName}_{Company}_Resume.docx``, strict ASCII.

    Both components go through :func:`_slug_ascii` so a non-ASCII company
    name (``"Nestlé"``) cannot corrupt the Content-Disposition header.
    """
    return f"{_slug_ascii(full_name)}_{_slug_ascii(company)}_Resume.docx"


def extract_docx_plaintext(path: Path | str) -> str:
    """Read any DOCX and return its paragraphs joined by single newlines.

    Distinct from :func:`extract_cover_letter_plaintext` (which uses a
    blank-line join suitable for plain-text email bodies). This helper
    is used by the Plan 05-04 pipeline to recover the tailored resume
    as plain text for the low-confidence holdout keyword-coverage
    check without loading the DOCX twice. Empty / whitespace-only
    paragraphs are dropped.
    """
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())


def extract_cover_letter_plaintext(cover_letter_path: Path | str) -> str:
    """Read the cover letter DOCX and return paragraph text joined by blank lines.

    Phase 4 only persists the cover letter as a DOCX at
    ``TailoringRecord.cover_letter_path``. There is no plaintext field.
    Re-extract via python-docx and join with ``"\\n\\n"`` — a single newline
    collapses paragraph breaks when placed into an email body (research
    pitfall 10). Empty / whitespace-only paragraphs are dropped so a stray
    blank line in the DOCX does not double-space the email.
    """
    from docx import Document

    doc = Document(str(cover_letter_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n\n".join(paragraphs)


def resolve_recipient_email(job_description: str) -> str | None:
    """Regex the first non-noreply email out of a job description.

    Returns ``None`` if no suitable address is found. Callers should flip
    ``Job.status = 'needs_info'`` on ``None`` (research open question #1).
    Noreply detection is prefix-based on the local part: ``noreply@``,
    ``no-reply@``, ``donotreply@``, ``do-not-reply@``, ``notifications@``.
    """
    if not job_description:
        return None
    for match in _EMAIL_RE.finditer(job_description):
        addr = match.group(0)
        local = addr.split("@", 1)[0].lower()
        if any(local.startswith(p) for p in _NOREPLY_PREFIXES):
            continue
        return addr
    return None


def build_email_message(
    *,
    from_addr: str,
    to_addr: str,
    subject: str,
    body_text: str,
    attachment_path: Path | str,
    attachment_filename: str,
) -> EmailMessage:
    """Construct an :class:`EmailMessage` with a single DOCX attachment.

    Uses the modern stdlib API (:meth:`EmailMessage.add_attachment`) which
    handles Base64 transfer encoding and RFC 2231 filename encoding for us.
    Do NOT use ``MIMEMultipart`` / ``MIMEBase`` (legacy, header-encoding bugs
    — see Phase 5 research §Don't Hand-Roll).
    """
    msg = EmailMessage()
    msg["From"] = from_addr
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.set_content(body_text)  # plain text, preserves newlines
    data = Path(attachment_path).read_bytes()
    msg.add_attachment(
        data,
        maintype=_DOCX_MAINTYPE,
        subtype=_DOCX_SUBTYPE,
        filename=attachment_filename,
    )
    return msg


__all__ = [
    "build_subject",
    "build_attachment_filename",
    "extract_docx_plaintext",
    "extract_cover_letter_plaintext",
    "resolve_recipient_email",
    "build_email_message",
]
