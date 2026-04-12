"""Format-preserving DOCX writer for tailored resumes and cover letters.

The tailored resume pipeline produces structured JSON sections (from the
LLM engine in ``app.tailoring.engine``) that must be written back into a
copy of the base resume DOCX without disturbing the user's original
fonts, bullet styles, bold/italic runs, spacing, or margins.

This module owns three responsibilities:

1. **build_tailored_docx** - Copy the base resume, walk paragraphs, and
   replace content section-by-section at the *run level* so formatting
   is preserved verbatim.  Heading paragraphs are never modified
   (headings are locked per the phase CONTEXT spec).
2. **build_cover_letter_docx** - Generate a clean business-letter DOCX
   from a list of paragraphs, matching the base resume's font family
   when detectable.
3. **check_ats_friendly / compute_keyword_coverage** - Post-generation
   ATS checks: tables present, non-standard fonts, keyword coverage
   ratio.

The hardest correctness constraint (per research Pitfall 1) is that we
must never use ``paragraph.text = ...`` when runs exist.  The setter
replaces all runs with a single plain run, losing every character
format.  All mutation goes through
``replace_paragraph_text_preserving_format``, which keeps the first
run's formatting as a template and drops the rest.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Any, Optional

import structlog
from docx import Document
from docx.shared import Pt

log = structlog.get_logger(__name__)


# -- Tunable constants --------------------------------------------------------

# Fonts generally considered safe by ATS systems.
_ATS_SAFE_FONTS: frozenset[str] = frozenset(
    {
        "Arial",
        "Calibri",
        "Times New Roman",
        "Helvetica",
        "Georgia",
        "Garamond",
        "Cambria",
        "Verdana",
    }
)

# Keywords used for fuzzy heading matching when Claude's section heading
# does not exactly match the base resume's heading (Pitfall 5).
_HEADING_KEYWORDS: tuple[str, ...] = (
    "experience",
    "education",
    "skills",
    "summary",
    "objective",
    "projects",
)


# -- Run-level replacement ---------------------------------------------------


def replace_paragraph_text_preserving_format(paragraph: Any, new_text: str) -> None:
    """Replace paragraph text while preserving run-level formatting.

    Strategy:

    - If the paragraph has no runs, fall back to the ``.text`` setter
      (there is no formatting to preserve).
    - Otherwise, keep the first run as the formatting template, remove
      every subsequent run via its lxml parent, and set the first run's
      text to the new content.

    This is the *only* path used to mutate existing paragraphs; the
    ``.text`` setter is never called when runs are present.
    """
    if not paragraph.runs:
        paragraph.text = new_text
        return

    first_run = paragraph.runs[0]
    for run in list(paragraph.runs[1:]):
        element = run._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    first_run.text = new_text


def _clear_paragraph(paragraph: Any) -> None:
    """Empty a paragraph's text while keeping its style handle.

    Used when the tailored section has fewer bullets than the base
    resume; the excess base paragraphs are cleared instead of deleted
    (deletion would shift surrounding spacing and re-flow the layout).
    """
    replace_paragraph_text_preserving_format(paragraph, "")


# -- Heading matching --------------------------------------------------------


def _match_heading(docx_heading: str, json_heading: str) -> bool:
    """Return True if a DOCX heading matches a JSON section heading.

    Matches in this order:

    1. Case-insensitive exact match after whitespace collapse.
    2. Keyword-based fuzzy match: if both headings contain the same
       canonical keyword (experience / education / skills / summary /
       objective / projects), they match.  This handles the common
       "Work Experience" vs "Professional Experience" divergence that
       research Pitfall 5 warns about.
    """
    if not docx_heading or not json_heading:
        return False

    a = re.sub(r"\s+", " ", docx_heading.strip().lower())
    b = re.sub(r"\s+", " ", json_heading.strip().lower())
    if a == b:
        return True

    for keyword in _HEADING_KEYWORDS:
        if keyword in a and keyword in b:
            return True

    return False


def _build_section_map(tailored_sections: dict) -> dict[str, dict]:
    """Index tailored sections by lowercased heading for O(1) lookup."""
    result: dict[str, dict] = {}
    for section in tailored_sections.get("sections", []) or []:
        heading = str(section.get("heading", "")).strip()
        if heading:
            result[heading.lower()] = section
    return result


def _find_tailored_section(
    docx_heading: str, section_map: dict[str, dict]
) -> Optional[dict]:
    """Look up a tailored section for a DOCX heading, falling back to fuzzy."""
    key = docx_heading.strip().lower()
    if key in section_map:
        return section_map[key]

    for json_heading, section in section_map.items():
        if _match_heading(docx_heading, json_heading):
            return section

    return None


# -- Section replacement primitives ------------------------------------------


def _replace_simple_section(
    paragraphs: list[Any],
    new_content: list[str],
) -> int:
    """Replace a run of non-heading paragraphs with new bullet strings.

    - Pairs tailored entries with base paragraphs in order.
    - If the tailored list is shorter, extra base paragraphs are
      cleared (see ``_clear_paragraph``).
    - If the tailored list is longer, the excess items are dropped
      with a warning log — adding paragraphs at arbitrary positions
      would require cloning style XML and tends to break spacing.

    Returns the number of paragraphs actually replaced.
    """
    replaced = 0
    for i, para in enumerate(paragraphs):
        if i < len(new_content):
            replace_paragraph_text_preserving_format(para, new_content[i])
            replaced += 1
        else:
            _clear_paragraph(para)

    if len(new_content) > len(paragraphs):
        overflow = len(new_content) - len(paragraphs)
        log.warning(
            "docx_writer.section_overflow",
            base_count=len(paragraphs),
            tailored_count=len(new_content),
            dropped=overflow,
        )

    return replaced


def _replace_experience_subsections(
    doc_paragraphs: list[Any],
    subsections: list[dict],
) -> int:
    """Replace work-experience subsections matched by locked company name.

    Work experience is structured as:
        Company Name (locked)
        Job Title / Dates (locked)
        - bullet 1
        - bullet 2

    We match base subsections to tailored subsections by the "company"
    field (case-insensitive contains match) and then replace only the
    bullet paragraphs — never the company / title / dates paragraphs,
    which are locked per CONTEXT.md.

    ``doc_paragraphs`` is the slice of doc.paragraphs that falls under
    the Work Experience heading.  Returns the number of bullet
    replacements performed.
    """
    if not subsections:
        return 0

    # Build a simple walking pointer over the base paragraphs.  We treat
    # any paragraph whose text is non-empty as a candidate; the first
    # paragraph of each subsection is expected to contain the company
    # name.  Bullets are everything up to the next paragraph whose text
    # matches another tailored company name.
    replaced = 0

    # Index tailored subsections by lowercased company token.
    by_company: list[tuple[str, dict]] = []
    for sub in subsections:
        company = str(sub.get("company", "")).strip().lower()
        if company:
            by_company.append((company, sub))

    if not by_company:
        return 0

    i = 0
    while i < len(doc_paragraphs):
        para = doc_paragraphs[i]
        text = para.text.strip()
        if not text:
            i += 1
            continue

        # Find a tailored subsection whose company token appears in
        # this paragraph's text.
        matched: Optional[dict] = None
        lowered = text.lower()
        for company, sub in by_company:
            if company and company in lowered:
                matched = sub
                break

        if matched is None:
            i += 1
            continue

        # Skip the company / title / dates rows.  Heuristic: skip
        # paragraphs until we hit one that looks like a bullet
        # (starts with a dash/bullet marker OR is followed by more
        # non-empty paragraphs).  We fall back to a fixed skip of 2
        # lines (company + title-dates) which is the typical layout.
        i += 1
        # Skip up to 2 more non-empty paragraphs that are likely
        # title / date lines (locked).
        skipped = 0
        while (
            i < len(doc_paragraphs)
            and skipped < 2
            and doc_paragraphs[i].text.strip()
            and not _looks_like_bullet(doc_paragraphs[i].text.strip())
        ):
            i += 1
            skipped += 1

        # Collect bullet paragraphs for this subsection: consecutive
        # non-empty paragraphs until we hit another company name or
        # run out of paragraphs.
        bullet_paras: list[Any] = []
        while i < len(doc_paragraphs):
            p = doc_paragraphs[i]
            ptext = p.text.strip()
            if not ptext:
                i += 1
                continue
            # Stop if this paragraph matches another tailored company.
            plower = ptext.lower()
            if any(c and c in plower for c, _ in by_company if c != matched.get("company", "").lower()):
                break
            bullet_paras.append(p)
            i += 1

        new_bullets = [str(b) for b in matched.get("bullets", []) or []]
        replaced += _replace_simple_section(bullet_paras, new_bullets)

    return replaced


def _looks_like_bullet(text: str) -> bool:
    """Heuristic: does this paragraph look like a bullet item?"""
    if not text:
        return False
    if text[0] in "-\u2022\u00b7*\u25cf\u25a0\u25e6":
        return True
    # "Company | dates" style lines usually contain a pipe or en-dash.
    return False


# -- Main writer -------------------------------------------------------------


def build_tailored_docx(
    base_resume_path: Path,
    tailored_sections: dict,
    output_path: Path,
) -> Path:
    """Write a tailored DOCX by modifying a copy of the base resume.

    Arguments:
        base_resume_path: Path to the user's base resume DOCX.
        tailored_sections: Structured JSON produced by the tailoring
            engine.  Expected shape::

                {
                    "sections": [
                        {"heading": "Summary", "content": ["..."]},
                        {"heading": "Experience", "subsections": [
                            {"company": "...", "title": "...",
                             "dates": "...", "bullets": ["..."]},
                        ]},
                    ],
                    "skills": ["Python", "FastAPI", ...],
                }
        output_path: Where to write the tailored DOCX.  Parent dirs
            will be created.

    Returns ``output_path``.

    The base resume is copied with ``shutil.copy2`` so the original
    file is never opened for writing.  Every heading in the copied
    document is walked; if the heading matches a tailored section, the
    section's content paragraphs are replaced in-place.  Headings that
    do not match any tailored section are left untouched.
    """
    base_resume_path = Path(base_resume_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    shutil.copy2(base_resume_path, output_path)
    doc = Document(str(output_path))

    section_map = _build_section_map(tailored_sections)
    if not section_map:
        log.warning("docx_writer.no_sections", path=str(base_resume_path))
        doc.save(str(output_path))
        return output_path

    # Group paragraphs by heading.
    groups: list[tuple[Optional[str], list[Any]]] = []
    current_heading: Optional[str] = None
    current_list: list[Any] = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            if current_heading is not None or current_list:
                groups.append((current_heading, current_list))
            current_heading = para.text.strip()
            current_list = []
        else:
            current_list.append(para)

    if current_heading is not None or current_list:
        groups.append((current_heading, current_list))

    total_replaced = 0
    matched_sections = 0

    for heading, paragraphs in groups:
        if heading is None:
            continue
        section = _find_tailored_section(heading, section_map)
        if section is None:
            continue

        matched_sections += 1
        subsections = section.get("subsections")
        content = section.get("content")

        if subsections:
            total_replaced += _replace_experience_subsections(paragraphs, subsections)
        elif isinstance(content, list):
            non_empty = [p for p in paragraphs if p.text.strip()]
            total_replaced += _replace_simple_section(
                non_empty, [str(c) for c in content]
            )
        elif isinstance(content, str):
            # Single-paragraph content (e.g. Summary).
            non_empty = [p for p in paragraphs if p.text.strip()]
            if non_empty:
                replace_paragraph_text_preserving_format(non_empty[0], content)
                total_replaced += 1

    # Skills block — may appear at top-level of the JSON rather than in
    # ``sections``, per the schema in 04-RESEARCH.md.
    skills = tailored_sections.get("skills")
    if isinstance(skills, list) and skills:
        skills_text = ", ".join(str(s) for s in skills)
        for heading, paragraphs in groups:
            if heading and _match_heading(heading, "Skills"):
                non_empty = [p for p in paragraphs if p.text.strip()]
                if non_empty:
                    replace_paragraph_text_preserving_format(non_empty[0], skills_text)
                    total_replaced += 1
                break

    log.info(
        "docx_writer.replaced",
        base=str(base_resume_path),
        output=str(output_path),
        matched_sections=matched_sections,
        paragraphs_replaced=total_replaced,
    )

    doc.save(str(output_path))
    return output_path


# -- Cover letter ------------------------------------------------------------


def _detect_base_font(base_resume_path: Optional[Path]) -> str:
    """Return the base resume's first-run font name, or 'Calibri'.

    Used so the cover letter visually matches the tailored resume.
    Any error reading the base resume falls back to the default.
    """
    if base_resume_path is None:
        return "Calibri"
    try:
        doc = Document(str(base_resume_path))
        for para in doc.paragraphs:
            if para.runs:
                name = para.runs[0].font.name
                if name:
                    return str(name)
    except Exception as exc:  # pragma: no cover - defensive
        log.debug("docx_writer.font_detect_failed", error=str(exc))
    return "Calibri"


def build_cover_letter_docx(
    paragraphs: list[str],
    output_path: Path,
    base_resume_path: Optional[Path] = None,
) -> Path:
    """Generate a clean business-letter DOCX from ``paragraphs``.

    Each string becomes a single paragraph with 11pt body text and
    standard business-letter spacing.  The default font is pulled from
    the base resume's first run when ``base_resume_path`` is provided,
    so cover letters visually match the tailored resume.

    Returns ``output_path``.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    font_name = _detect_base_font(base_resume_path)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(11)

    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(10)

    doc.save(str(output_path))
    log.info(
        "docx_writer.cover_letter_written",
        output=str(output_path),
        paragraphs=len(paragraphs),
        font=font_name,
    )
    return output_path


# -- ATS checks --------------------------------------------------------------


def check_ats_friendly(docx_path: Path) -> dict:
    """Inspect a DOCX for ATS-incompatible features.

    Returns a dict with::

        {
            "has_tables": bool,
            "non_standard_fonts": list[str],  # unique, sorted
            "keyword_coverage": None,  # computed separately
        }

    The ``keyword_coverage`` field is always ``None`` here because it
    depends on the job description, which this function does not see.
    Callers that have the job text should merge in the value returned
    by :func:`compute_keyword_coverage`.
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))

    has_tables = len(doc.tables) > 0

    seen_fonts: set[str] = set()
    for para in doc.paragraphs:
        for run in para.runs:
            name = run.font.name
            if name:
                seen_fonts.add(str(name))

    non_standard = sorted(f for f in seen_fonts if f not in _ATS_SAFE_FONTS)

    result = {
        "has_tables": has_tables,
        "non_standard_fonts": non_standard,
        "keyword_coverage": None,
    }
    log.info(
        "docx_writer.ats_check",
        path=str(docx_path),
        has_tables=has_tables,
        non_standard_fonts=non_standard,
    )
    return result


def compute_keyword_coverage(tailored_text: str, job_description: str) -> float:
    """Return the fraction of job-description keywords present in ``tailored_text``.

    Keywords are extracted from the job description as word tokens
    longer than 3 characters; stopwords aren't stripped because the
    tailored text will organically cover most common English words.
    Matching is case-insensitive.

    Returns ``0.0`` when the job description has no usable tokens.
    """
    if not job_description:
        return 0.0

    tokens = {w.lower() for w in re.findall(r"[A-Za-z][A-Za-z+\-/]*", job_description)}
    tokens = {t for t in tokens if len(t) > 3}
    if not tokens:
        return 0.0

    lowered = tailored_text.lower()
    matches = sum(1 for t in tokens if t in lowered)
    return round(matches / len(tokens), 4)


__all__ = [
    "replace_paragraph_text_preserving_format",
    "build_tailored_docx",
    "build_cover_letter_docx",
    "check_ats_friendly",
    "compute_keyword_coverage",
]
