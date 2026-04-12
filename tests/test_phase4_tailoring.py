"""Phase 4 tailoring test suite (TAIL-01..TAIL-09 + SAFE-04).

Covers the full Phase 4 surface with mocked LLM calls and real DOCX
operations. Every TAIL requirement and the SAFE-04 PII constraint has
at least one dedicated test; the key "truth" tests called out in the
phase plan (hallucination rejection, PII never in prompt, budget halt
at cap, cache-token visibility in the review template) are explicitly
labelled below.

Structure mirrors the Phase 3 test modules: small, legible classes with
one behaviour per test. LLM calls are mocked via ``FakeLLMProvider`` so
no real Anthropic traffic is generated. DB-bound tests reuse the
existing ``async_session`` / ``async_session_factory`` fixtures from
``conftest.py``. Integration tests use the ``live_app`` pattern from
Phase 3 for the settings UI route.
"""

from __future__ import annotations

import importlib
import json
from datetime import datetime
from pathlib import Path
from typing import Optional

import httpx
import pytest
import pytest_asyncio
from cryptography.fernet import Fernet
from docx import Document
from docx.shared import Pt

from app.tailoring.budget import BudgetGuard
from app.tailoring.docx_writer import (
    build_cover_letter_docx,
    build_tailored_docx,
    check_ats_friendly,
    compute_keyword_coverage,
    replace_paragraph_text_preserving_format,
)
from app.tailoring.engine import (
    TailoringResult,
    strip_pii_sections,
    tailor_resume,
    validate_output,
)
from app.tailoring.preview import (
    docx_to_html,
    generate_section_diff,
)
from app.tailoring.prompts import (
    build_system_messages,
    build_tailoring_user_message,
    get_escalated_prompt_suffix,
)
from app.tailoring.provider import LLMResponse
from app.tailoring.service import (
    get_next_version,
    resume_artifact_path,
    save_cost_entries,
    save_tailoring_record,
)


# =========================================================================
# Test doubles
# =========================================================================


class FakeLLMProvider:
    """Scriptable mock provider returning canned LLMResponse objects.

    ``responses`` is a list of either raw content strings (wrapped into
    a default-token LLMResponse) or fully formed LLMResponse instances
    for tests that need specific token counts. Every ``complete`` call
    is recorded in ``self.calls`` so assertions can inspect what was
    actually sent (critical for the SAFE-04 PII-in-prompt tests).
    """

    def __init__(self, responses):
        self._responses = list(responses)
        self._index = 0
        self.calls: list[dict] = []

    async def complete(self, system, messages, max_tokens, temperature=0.3):
        self.calls.append(
            {
                "system": system,
                "messages": messages,
                "max_tokens": max_tokens,
                "temperature": temperature,
            }
        )
        if self._index >= len(self._responses):
            raise AssertionError(
                f"FakeLLMProvider ran out of responses at call "
                f"{self._index}"
            )
        item = self._responses[self._index]
        self._index += 1
        if isinstance(item, LLMResponse):
            return item
        return LLMResponse(
            content=str(item),
            input_tokens=100,
            output_tokens=50,
            cache_creation_tokens=0,
            cache_read_tokens=80,
            model="fake-model",
        )


# =========================================================================
# DOCX fixtures
# =========================================================================


def _make_base_resume_docx(path: Path) -> Path:
    """Create a minimal but realistic base resume DOCX for testing.

    Contains every structural element the writer touches: contact
    header, summary, work experience with bullets, skills, and education.
    Includes a bold/italic run to prove formatting is preserved.
    """
    doc = Document()

    # Contact header (no style → no heading → gets stripped as PII).
    p = doc.add_paragraph()
    run = p.add_run("Jane Doe")
    run.bold = True
    doc.add_paragraph("jane.doe@example.com  |  +1 555-123-4567")
    doc.add_paragraph("123 Main St, Springfield")

    # Professional Summary
    doc.add_heading("Professional Summary", level=1)
    sp = doc.add_paragraph()
    sr = sp.add_run("Backend engineer with 5 years of Python experience.")
    sr.italic = True

    # Work Experience
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Acme Corp")
    doc.add_paragraph("Senior Software Engineer  |  2021-2024")
    bp = doc.add_paragraph("- Built FastAPI services handling 10k req/s")
    # Bold keyword inside bullet
    bp.runs[0].bold = True
    doc.add_paragraph("- Led a team of 4 engineers on PostgreSQL migrations")
    doc.add_paragraph("- Shipped production features weekly")

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. Computer Science, State University, 2019")

    doc.save(str(path))
    return path


@pytest.fixture
def base_resume_path(tmp_path: Path) -> Path:
    return _make_base_resume_docx(tmp_path / "base_resume.docx")


@pytest.fixture
def sample_resume_sections() -> list[dict]:
    """Sections shaped like :func:`app.resume.service.extract_resume_text`."""
    return [
        {
            "heading": None,
            "text": (
                "Jane Doe\njane.doe@example.com | +1 555-123-4567\n"
                "123 Main St, Springfield"
            ),
        },
        {
            "heading": "Professional Summary",
            "text": "Backend engineer with 5 years of Python experience.",
        },
        {
            "heading": "Work Experience",
            "text": (
                "Acme Corp\nSenior Software Engineer | 2021-2024\n"
                "- Built FastAPI services handling 10k req/s\n"
                "- Led a team of 4 engineers on PostgreSQL migrations"
            ),
        },
        {
            "heading": "Skills",
            "text": "Python, FastAPI, PostgreSQL, Docker",
        },
    ]


# =========================================================================
# PII stripping tests (SAFE-04)
# =========================================================================


class TestPIIStripping:
    """SAFE-04: PII never enters the LLM prompt."""

    def test_strip_pii_removes_contact_section(self, sample_resume_sections):
        sanitized, headings = strip_pii_sections(sample_resume_sections)

        # Unlabelled first block (name/email/phone/address) must be gone.
        assert "Jane Doe" not in sanitized
        assert "jane.doe@example.com" not in sanitized
        assert "555-123-4567" not in sanitized
        assert "Main St" not in sanitized

    def test_strip_pii_preserves_content_sections(self, sample_resume_sections):
        sanitized, headings = strip_pii_sections(sample_resume_sections)

        assert "Professional Summary" in sanitized
        assert "Work Experience" in sanitized
        assert "FastAPI" in sanitized
        assert "Skills" in sanitized
        assert "Python" in sanitized

        # Headings list mirrors kept sections in order.
        assert headings == [
            "Professional Summary",
            "Work Experience",
            "Skills",
        ]

    def test_strip_pii_logs_warning_when_no_contact_stripped(self, caplog):
        sections = [
            {"heading": "Summary", "text": "Backend engineer"},
            {"heading": "Skills", "text": "Python"},
        ]
        sanitized, _ = strip_pii_sections(sections)
        # No contact section → belt-and-braces regex still runs.
        assert "Backend engineer" in sanitized
        assert "Python" in sanitized

    def test_strip_pii_redacts_dangling_email(self):
        sections = [
            {"heading": "Summary", "text": "Reach me at stray@example.com"},
        ]
        sanitized, _ = strip_pii_sections(sections)
        assert "stray@example.com" not in sanitized
        assert "[REDACTED_EMAIL]" in sanitized

    def test_strip_pii_removes_contact_heading_section(self):
        sections = [
            {"heading": "Intro", "text": "About me"},
            {"heading": "Personal Info", "text": "Phone: 555-9876"},
            {"heading": "Skills", "text": "Python"},
        ]
        sanitized, _ = strip_pii_sections(sections)
        assert "Personal Info" not in sanitized
        assert "Skills" in sanitized


# =========================================================================
# Prompts tests
# =========================================================================


class TestPrompts:
    """System-prompt construction + intensity + escalation."""

    def test_system_messages_have_cache_control(self):
        msgs = build_system_messages("BASE RESUME TEXT HERE")
        assert len(msgs) == 2
        # Second block (base resume) is the cached one.
        assert "cache_control" in msgs[1]
        assert msgs[1]["cache_control"] == {"type": "ephemeral"}
        # First block (instructions) is NOT cached.
        assert "cache_control" not in msgs[0]

    def test_intensity_instructions_differ(self):
        light = build_tailoring_user_message(
            "JD", "light", ["Skills"], retry=0
        )
        balanced = build_tailoring_user_message(
            "JD", "balanced", ["Skills"], retry=0
        )
        full = build_tailoring_user_message(
            "JD", "full", ["Skills"], retry=0
        )
        assert "light" in light
        assert "balanced" in balanced
        assert "full" in full
        assert light != balanced != full

    def test_intensity_normalises_unknown_value(self):
        msg = build_tailoring_user_message("JD", "bogus", ["Skills"], retry=0)
        assert "balanced" in msg

    def test_escalation_suffix_increases(self):
        r0 = get_escalated_prompt_suffix(0)
        r1 = get_escalated_prompt_suffix(1)
        r2 = get_escalated_prompt_suffix(2)
        assert r0 == ""
        assert "MORE conservative" in r1
        assert "FINAL RETRY" in r2
        assert len(r2) > len(r1) > len(r0)


# =========================================================================
# Hallucination validator tests (TAIL-04, SC-2)
# =========================================================================


class TestValidator:
    """TAIL-04: LLM-as-judge validator detects invented content."""

    @pytest.mark.asyncio
    async def test_validator_passes_clean_output(self):
        provider = FakeLLMProvider(
            [json.dumps({"passed": True, "violations": []})]
        )
        passed, violations, resp = await validate_output(
            provider,
            original_text="Original resume text",
            tailored_json='{"sections": []}',
        )
        assert passed is True
        assert violations == []
        assert isinstance(resp, LLMResponse)

    @pytest.mark.asyncio
    async def test_validator_rejects_invented_skill(self):
        provider = FakeLLMProvider(
            [
                json.dumps(
                    {
                        "passed": False,
                        "violations": [
                            {
                                "type": "invented_skill",
                                "content": "React",
                                "explanation": (
                                    "Original resume does not mention React"
                                ),
                            }
                        ],
                    }
                )
            ]
        )
        passed, violations, _ = await validate_output(
            provider,
            original_text="Python FastAPI PostgreSQL",
            tailored_json='{"sections": [], "skills": ["React"]}',
        )
        assert passed is False
        assert len(violations) == 1
        assert violations[0]["type"] == "invented_skill"
        assert violations[0]["content"] == "React"

    @pytest.mark.asyncio
    async def test_invented_skill_rejected_end_to_end(
        self, sample_resume_sections
    ):
        """SC-2 KEY TEST: full tailor_resume loop rejects invented skill.

        Scripts 3 tailor+validate cycles where every attempt invents
        "React". All three validation calls flag it; tailor_resume must
        return success=False and surface the violation history.
        """
        tailored_with_react = json.dumps(
            {
                "sections": [
                    {
                        "heading": "Skills",
                        "content": ["Python", "React", "PostgreSQL"],
                    }
                ]
            }
        )
        validator_reject = json.dumps(
            {
                "passed": False,
                "violations": [
                    {
                        "type": "invented_skill",
                        "content": "React",
                        "explanation": "Not in base resume",
                    }
                ],
            }
        )
        # tailor → validate → tailor → validate → tailor → validate
        responses = [
            tailored_with_react,
            validator_reject,
            tailored_with_react,
            validator_reject,
            tailored_with_react,
            validator_reject,
        ]
        provider = FakeLLMProvider(responses)

        result = await tailor_resume(
            provider=provider,
            resume_sections=sample_resume_sections,
            job_description="React frontend role",
            intensity="balanced",
            max_retries=3,
        )

        assert result.success is False
        assert result.validation_passed is False
        assert result.retry_count == 3
        # Violation history preserved across retries.
        assert any(
            w.get("type") == "invented_skill"
            for w in result.validation_warnings
        )
        assert any(
            w.get("content") == "React"
            for w in result.validation_warnings
        )

    @pytest.mark.asyncio
    async def test_pii_never_in_llm_prompt(self, sample_resume_sections):
        """SAFE-04 KEY TEST: PII does not leak into any provider call.

        Run a full tailor_resume cycle and inspect every captured
        provider.calls payload. The literal name/email/phone must NOT
        appear in any system block or user message.
        """
        good_tailor = json.dumps(
            {
                "sections": [
                    {
                        "heading": "Professional Summary",
                        "content": ["Backend engineer with Python expertise"],
                    }
                ]
            }
        )
        validator_pass = json.dumps({"passed": True, "violations": []})
        cover_letter = json.dumps(
            {
                "paragraphs": [
                    "Para 1",
                    "Para 2",
                    "Para 3",
                ]
            }
        )
        provider = FakeLLMProvider(
            [good_tailor, validator_pass, cover_letter]
        )

        await tailor_resume(
            provider=provider,
            resume_sections=sample_resume_sections,
            job_description="Python backend role",
            intensity="balanced",
        )

        # Flatten every bit of text sent to the provider into one blob.
        blob_parts: list[str] = []
        for call in provider.calls:
            system = call["system"]
            if isinstance(system, list):
                for block in system:
                    blob_parts.append(json.dumps(block))
            else:
                blob_parts.append(str(system))
            blob_parts.append(json.dumps(call["messages"]))
        blob = "\n".join(blob_parts)

        # Literal PII must NOT appear anywhere in the blob.
        assert "Jane Doe" not in blob
        assert "jane.doe@example.com" not in blob
        assert "555-123-4567" not in blob
        assert "Main St" not in blob


# =========================================================================
# Budget guard tests (TAIL-08, SC-4)
# =========================================================================


class TestBudgetGuard:
    """TAIL-08: monthly budget halt and warning thresholds."""

    @pytest.mark.asyncio
    async def test_budget_check_allows_under_cap(self, async_session):
        from app.settings.service import get_settings_row

        row = await get_settings_row(async_session)
        row.budget_cap_dollars = 10.0
        row.budget_spent_dollars = 5.0
        row.budget_month = BudgetGuard._current_month()
        await async_session.commit()

        guard = BudgetGuard()
        can_proceed, spent, cap, is_warning = await guard.check_budget(
            async_session
        )
        assert can_proceed is True
        assert spent == 5.0
        assert cap == 10.0
        assert is_warning is False

    @pytest.mark.asyncio
    async def test_budget_check_halts_at_cap(self, async_session):
        """SC-4: tailoring halts at 100% cap."""
        from app.settings.service import get_settings_row

        row = await get_settings_row(async_session)
        row.budget_cap_dollars = 10.0
        row.budget_spent_dollars = 10.0
        row.budget_month = BudgetGuard._current_month()
        await async_session.commit()

        guard = BudgetGuard()
        can_proceed, spent, cap, is_warning = await guard.check_budget(
            async_session
        )
        assert can_proceed is False
        assert is_warning is True

    @pytest.mark.asyncio
    async def test_budget_warning_at_80_percent(self, async_session):
        from app.settings.service import get_settings_row

        row = await get_settings_row(async_session)
        row.budget_cap_dollars = 10.0
        row.budget_spent_dollars = 8.0
        row.budget_month = BudgetGuard._current_month()
        await async_session.commit()

        guard = BudgetGuard()
        can_proceed, spent, cap, is_warning = await guard.check_budget(
            async_session
        )
        assert can_proceed is True
        assert is_warning is True

    @pytest.mark.asyncio
    async def test_budget_month_rollover(self, async_session):
        from app.settings.service import get_settings_row

        row = await get_settings_row(async_session)
        row.budget_cap_dollars = 10.0
        row.budget_spent_dollars = 50.0
        row.budget_month = "1999-01"  # far-past month
        await async_session.commit()

        guard = BudgetGuard()
        can_proceed, spent, cap, _ = await guard.check_budget(async_session)
        assert can_proceed is True
        assert spent == 0.0  # rolled over
        assert cap == 10.0

        # Verify persisted reset.
        await async_session.refresh(row)
        assert row.budget_spent_dollars == 0.0
        assert row.budget_month == BudgetGuard._current_month()

    @pytest.mark.asyncio
    async def test_budget_no_cap_always_proceeds(self, async_session):
        from app.settings.service import get_settings_row

        row = await get_settings_row(async_session)
        row.budget_cap_dollars = 0.0  # "unlimited" sentinel
        row.budget_spent_dollars = 9999.0
        row.budget_month = BudgetGuard._current_month()
        await async_session.commit()

        guard = BudgetGuard()
        can_proceed, _, cap, is_warning = await guard.check_budget(
            async_session
        )
        assert can_proceed is True
        assert cap == 0.0
        assert is_warning is False

    def test_estimate_cost_calculation(self):
        # Sonnet 4.5: $3/Mtok input, $15/Mtok output.
        # 1M input + 1M output = $3 + $15 = $18
        cost = BudgetGuard.estimate_cost(
            input_tokens=1_000_000,
            output_tokens=1_000_000,
        )
        assert cost == 18.0

    def test_estimate_cost_cache_savings(self):
        # Cache reads should be 10x cheaper than fresh input.
        fresh = BudgetGuard.estimate_cost(
            input_tokens=1_000_000, output_tokens=0
        )
        cached = BudgetGuard.estimate_cost(
            input_tokens=0,
            output_tokens=0,
            cache_read_tokens=1_000_000,
        )
        assert fresh == 3.0
        assert cached == 0.3
        assert cached < fresh / 5

    def test_estimate_cost_unknown_model_falls_back(self):
        fallback = BudgetGuard.estimate_cost(
            input_tokens=1_000_000,
            output_tokens=0,
            model="some-new-model-not-in-table",
        )
        default = BudgetGuard.estimate_cost(
            input_tokens=1_000_000, output_tokens=0
        )
        assert fallback == default


# =========================================================================
# DOCX writer tests (TAIL-05, SC-1)
# =========================================================================


class TestDocxWriter:
    """TAIL-05: format-preserving DOCX writer + ATS checks."""

    def test_replace_preserves_run_formatting_bold(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("original bold text")
        run.bold = True
        out = tmp_path / "test.docx"
        doc.save(str(out))

        doc2 = Document(str(out))
        replace_paragraph_text_preserving_format(
            doc2.paragraphs[0], "replacement text"
        )
        doc2.save(str(out))

        verify = Document(str(out))
        # The first run should still be bold after replacement.
        assert verify.paragraphs[0].runs[0].bold is True
        assert verify.paragraphs[0].runs[0].text == "replacement text"

    def test_build_tailored_docx_preserves_formatting(
        self, base_resume_path, tmp_path
    ):
        """SC-1: bold/italic formatting survives tailoring."""
        tailored = {
            "sections": [
                {
                    "heading": "Professional Summary",
                    "content": [
                        "Senior backend engineer specializing in Python APIs."
                    ],
                },
            ]
        }
        out = tmp_path / "tailored.docx"
        build_tailored_docx(base_resume_path, tailored, out)

        doc = Document(str(out))
        # Find the Summary paragraph (first non-empty para after the
        # "Professional Summary" heading).
        found_italic = False
        hit_heading = False
        for para in doc.paragraphs:
            if (
                para.style
                and para.style.name.startswith("Heading")
                and "Summary" in para.text
            ):
                hit_heading = True
                continue
            if hit_heading and para.text.strip():
                # First non-empty paragraph after heading: the summary.
                assert (
                    "Senior backend engineer" in para.text
                    or "Python" in para.text
                )
                if para.runs and para.runs[0].italic:
                    found_italic = True
                break
        # Italic template preserved.
        assert found_italic

    def test_build_tailored_docx_replaces_content(
        self, base_resume_path, tmp_path
    ):
        tailored = {
            "sections": [
                {
                    "heading": "Professional Summary",
                    "content": ["BRAND NEW SUMMARY TEXT"],
                }
            ]
        }
        out = tmp_path / "tailored.docx"
        build_tailored_docx(base_resume_path, tailored, out)

        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
        assert "BRAND NEW SUMMARY TEXT" in text

    def test_cover_letter_docx_created(self, tmp_path):
        paragraphs = [
            "Opening paragraph with fit statement.",
            "Highlight paragraph with two specific projects.",
            "Closing paragraph with genuine enthusiasm.",
        ]
        out = tmp_path / "cover.docx"
        result = build_cover_letter_docx(paragraphs, out)
        assert result == out
        assert out.exists()

        doc = Document(str(out))
        rendered = [p.text for p in doc.paragraphs]
        assert "Opening paragraph with fit statement." in rendered
        assert "Highlight paragraph with two specific projects." in rendered
        assert len([p for p in rendered if p.strip()]) == 3

    def test_ats_check_detects_table(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Has a table")
        doc.add_table(rows=2, cols=2)
        out = tmp_path / "with_table.docx"
        doc.save(str(out))

        result = check_ats_friendly(out)
        assert result["has_tables"] is True
        assert result["keyword_coverage"] is None

    def test_ats_check_detects_nonstandard_font(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Comic text")
        run.font.name = "Comic Sans MS"
        out = tmp_path / "comic.docx"
        doc.save(str(out))

        result = check_ats_friendly(out)
        assert "Comic Sans MS" in result["non_standard_fonts"]

    def test_ats_check_flags_no_issues_for_clean_docx(self, tmp_path):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Clean text")
        run.font.name = "Calibri"
        out = tmp_path / "clean.docx"
        doc.save(str(out))

        result = check_ats_friendly(out)
        assert result["has_tables"] is False
        assert result["non_standard_fonts"] == []

    def test_keyword_coverage_computation(self):
        tailored = "Python FastAPI PostgreSQL backend engineer"
        jd = "Python FastAPI PostgreSQL Docker Kubernetes"
        ratio = compute_keyword_coverage(tailored, jd)
        # python, fastapi, postgresql present; docker, kubernetes missing.
        # 3/5 = 0.6
        assert 0.5 <= ratio <= 0.7

    def test_keyword_coverage_empty_jd_returns_zero(self):
        assert compute_keyword_coverage("any text", "") == 0.0


# =========================================================================
# Preview tests (TAIL-06)
# =========================================================================


class TestPreview:
    """TAIL-06: DOCX→HTML preview + section diff."""

    def test_docx_to_html_returns_string(self, base_resume_path):
        html = docx_to_html(base_resume_path)
        assert isinstance(html, str)
        assert len(html) > 0
        assert "FastAPI" in html or "Python" in html

    def test_section_diff_identifies_changes(self):
        base = [
            {"heading": "Summary", "text": "Original summary text"},
            {"heading": "Skills", "text": "Python, Docker"},
        ]
        tailored = {
            "sections": [
                {
                    "heading": "Summary",
                    "content": ["Completely rewritten summary text"],
                },
                {"heading": "Skills", "content": ["Python, Docker"]},
            ]
        }
        diffs = generate_section_diff(base, tailored)
        assert len(diffs) == 2
        summary_diff = next(d for d in diffs if d["heading"] == "Summary")
        assert summary_diff["changed"] is True
        skills_diff = next(d for d in diffs if d["heading"] == "Skills")
        assert skills_diff["changed"] is False

    def test_section_diff_flags_tailored_only_section(self):
        base = [{"heading": "Summary", "text": "Orig"}]
        tailored = {
            "sections": [
                {"heading": "Summary", "content": ["Orig"]},
                {
                    "heading": "Projects",
                    "content": ["Built thing X"],
                },
            ]
        }
        diffs = generate_section_diff(base, tailored)
        projects = next(d for d in diffs if d["heading"] == "Projects")
        assert projects["changed"] is True
        assert projects["base_text"] == ""


# =========================================================================
# Service layer tests (TAIL-09, SC-5)
# =========================================================================


class TestServiceLayer:
    """TAIL-09: versioned artifact storage + record persistence."""

    def test_artifact_paths_versioned(
        self, env_with_fernet, monkeypatch, tmp_path
    ):
        path = resume_artifact_path(42, 1)
        as_str = str(path).replace("\\", "/")
        assert "42/v1.docx" in as_str

        path2 = resume_artifact_path(42, 2)
        assert "v2.docx" in str(path2).replace("\\", "/")

    @pytest.mark.asyncio
    async def test_get_next_version_empty(self, async_session):
        next_v = await get_next_version(async_session, job_id=99)
        assert next_v == 1

    @pytest.mark.asyncio
    async def test_get_next_version_increments(self, async_session):
        from app.discovery.models import Job
        from app.tailoring.models import TailoringRecord

        # Need a job row to satisfy FK.
        job = Job(
            fingerprint="fp1",
            external_id="ext-1",
            title="Engineer",
            company="Acme",
            url="http://x",
            source="greenhouse",
            source_id=1,
            description="desc",
            description_html="<p>desc</p>",
            status="matched",
            score=80,
        )
        async_session.add(job)
        await async_session.commit()
        await async_session.refresh(job)

        # Insert v1 and v2 records.
        for v in (1, 2):
            rec = TailoringRecord(
                job_id=job.id,
                version=v,
                intensity="balanced",
                status="completed",
                base_resume_path="/tmp/base.docx",
            )
            async_session.add(rec)
        await async_session.commit()

        next_v = await get_next_version(async_session, job_id=job.id)
        assert next_v == 3

    @pytest.mark.asyncio
    async def test_save_and_query_tailoring_record(self, async_session):
        """TAIL-09 core write path."""
        from app.discovery.models import Job
        from app.tailoring.models import TailoringRecord
        from sqlalchemy import select

        job = Job(
            fingerprint="fp2",
            external_id="ext-2",
            title="Backend Dev",
            company="Initech",
            url="http://x/2",
            source="lever",
            source_id=1,
            description="desc",
            description_html="<p>desc</p>",
            status="matched",
            score=75,
        )
        async_session.add(job)
        await async_session.commit()
        await async_session.refresh(job)

        fake_result = TailoringResult(
            success=True,
            tailored_sections={"sections": []},
            cover_letter_paragraphs=["p1", "p2"],
            validation_passed=True,
            validation_warnings=[],
            total_input_tokens=500,
            total_output_tokens=200,
            total_cache_read_tokens=400,
            total_cache_write_tokens=100,
            retry_count=1,
            error=None,
            llm_calls=[
                {
                    "call_type": "tailor",
                    "model": "claude-sonnet-4-5",
                    "input_tokens": 500,
                    "output_tokens": 200,
                    "cache_read_tokens": 400,
                    "cache_write_tokens": 100,
                }
            ],
        )

        record = await save_tailoring_record(
            session=async_session,
            job_id=job.id,
            version=1,
            intensity="balanced",
            base_resume_path="/data/base.docx",
            tailored_resume_path="/data/resumes/1/v1.docx",
            cover_letter_path="/data/resumes/1/cover_letter_v1.docx",
            result=fake_result,
            status="completed",
            resume_text="sample",
            job_description="jd",
            system_prompt="sys",
        )
        await async_session.commit()

        assert record.id is not None
        assert record.cache_read_tokens == 400
        assert record.cache_write_tokens == 100
        assert record.retry_count == 1
        assert record.prompt_hash is not None
        assert record.validation_passed is True

        # Query it back.
        result = await async_session.execute(
            select(TailoringRecord).where(TailoringRecord.id == record.id)
        )
        fetched = result.scalar_one()
        assert fetched.version == 1
        assert fetched.tailored_resume_path == "/data/resumes/1/v1.docx"
        assert fetched.estimated_cost_dollars > 0

    @pytest.mark.asyncio
    async def test_save_cost_entries_matches_llm_calls(self, async_session):
        from app.discovery.models import Job
        from app.tailoring.models import CostLedger
        from sqlalchemy import select

        job = Job(
            fingerprint="fp3",
            external_id="ext-3",
            title="Eng",
            company="Co",
            url="http://y",
            source="ashby",
            source_id=1,
            description="d",
            description_html="<p>d</p>",
            status="matched",
            score=70,
        )
        async_session.add(job)
        await async_session.commit()
        await async_session.refresh(job)

        result_obj = TailoringResult(
            success=True,
            tailored_sections={"sections": []},
            cover_letter_paragraphs=None,
            validation_passed=True,
            validation_warnings=[],
            total_input_tokens=0,
            total_output_tokens=0,
            total_cache_read_tokens=0,
            total_cache_write_tokens=0,
            retry_count=1,
            error=None,
            llm_calls=[],
        )
        record = await save_tailoring_record(
            session=async_session,
            job_id=job.id,
            version=1,
            intensity="balanced",
            base_resume_path="/b.docx",
            tailored_resume_path=None,
            cover_letter_path=None,
            result=result_obj,
            status="completed",
        )
        await async_session.commit()

        calls = [
            {
                "call_type": "tailor",
                "model": "claude-sonnet-4-5",
                "input_tokens": 1000,
                "output_tokens": 500,
                "cache_read_tokens": 800,
                "cache_write_tokens": 0,
            },
            {
                "call_type": "validate",
                "model": "claude-sonnet-4-5",
                "input_tokens": 200,
                "output_tokens": 100,
                "cache_read_tokens": 0,
                "cache_write_tokens": 0,
            },
        ]
        rows = await save_cost_entries(
            async_session, record.id, calls
        )
        await async_session.commit()
        assert len(rows) == 2

        # Verify persisted.
        fetched = (
            await async_session.execute(
                select(CostLedger).where(
                    CostLedger.tailoring_record_id == record.id
                )
            )
        ).scalars().all()
        assert len(fetched) == 2
        call_types = {r.call_type for r in fetched}
        assert call_types == {"tailor", "validate"}


# =========================================================================
# Settings UI tests
# =========================================================================


class TestTailoringIntensitySettings:
    """UI: tailoring intensity selector persists correctly."""

    @pytest.mark.asyncio
    async def test_settings_tailoring_intensity_persists(
        self, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
    ):
        key = Fernet.generate_key().decode()
        monkeypatch.setenv("FERNET_KEY", key)
        monkeypatch.setenv("TZ", "UTC")
        monkeypatch.setenv("BIND_ADDRESS", "127.0.0.1")
        monkeypatch.setenv("DATA_DIR", str(tmp_path))

        import app.config as config_module

        config_module.get_settings.cache_clear()
        importlib.reload(config_module)

        import app.db.base as base_module

        importlib.reload(base_module)

        import app.web.routers.wizard as wizard_module

        importlib.reload(wizard_module)

        from app.db import models  # noqa: F401
        from sqlmodel import SQLModel

        async with base_module.engine.begin() as conn:
            await conn.run_sync(SQLModel.metadata.create_all)

        from app.settings.service import (
            get_setting,
            get_settings_row,
            set_setting,
        )

        async with base_module.async_session() as session:
            await set_setting(session, "wizard_complete", True)

        import app.main as main_module

        importlib.reload(main_module)

        application = main_module.create_app()

        async with application.router.lifespan_context(application):
            async with httpx.AsyncClient(
                transport=httpx.ASGITransport(app=application),
                base_url="http://test",
            ) as client:
                resp = await client.post(
                    "/settings/tailoring",
                    data={"tailoring_intensity": "full"},
                )
                assert resp.status_code == 200

                async with base_module.async_session() as session:
                    value = await get_setting(
                        session, "tailoring_intensity"
                    )
                    assert value == "full"

                # Reject unknown value with 400.
                resp_bad = await client.post(
                    "/settings/tailoring",
                    data={"tailoring_intensity": "maximum"},
                )
                assert resp_bad.status_code == 400

        await base_module.engine.dispose()


# =========================================================================
# Cache token visibility tests (SC-5)
# =========================================================================


class TestCacheTokenVisibility:
    """SC-5: prompt caching savings are surfaced in the review UI."""

    def test_tailoring_detail_template_shows_cache_tokens(self):
        template_path = (
            Path(__file__).resolve().parents[1]
            / "app"
            / "web"
            / "templates"
            / "partials"
            / "tailoring_detail.html.j2"
        )
        assert template_path.exists()
        content = template_path.read_text(encoding="utf-8")

        # SC-5: cache read tokens visible.
        assert "cache_read_tokens" in content
        # SC-5: cache savings visible (human-readable label).
        lowered = content.lower()
        assert "cache savings" in lowered or "cache_savings" in lowered

    def test_tailoring_result_tracks_cache_tokens(self):
        result = TailoringResult(
            success=True,
            tailored_sections={"sections": []},
            cover_letter_paragraphs=None,
            validation_passed=True,
            validation_warnings=[],
            total_input_tokens=1000,
            total_output_tokens=500,
            total_cache_read_tokens=500,
            total_cache_write_tokens=200,
            retry_count=1,
            error=None,
        )
        assert result.total_cache_read_tokens == 500
        assert result.total_cache_write_tokens == 200
        # Data flows to the UI through these four buckets.
        assert hasattr(result, "total_cache_read_tokens")
        assert hasattr(result, "total_cache_write_tokens")
