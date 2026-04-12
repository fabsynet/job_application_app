"""Phase 2 integration tests: resume upload and DOCX extraction (CONF-03).

Tests exercise the full HTTP -> service -> filesystem flow using a real
DOCX file created via python-docx in each test.
"""

from __future__ import annotations

import importlib
import io
from pathlib import Path

import httpx
import pytest
import pytest_asyncio
from cryptography.fernet import Fernet
from docx import Document


@pytest_asyncio.fixture
async def live_app(monkeypatch: pytest.MonkeyPatch, tmp_path: Path):
    key = Fernet.generate_key().decode()
    monkeypatch.setenv("FERNET_KEY", key)
    monkeypatch.setenv("TZ", "UTC")
    monkeypatch.setenv("BIND_ADDRESS", "127.0.0.1")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("LOG_LEVEL", "INFO")

    import app.config as config_module

    config_module.get_settings.cache_clear()
    importlib.reload(config_module)

    import app.db.base as base_module

    importlib.reload(base_module)

    import app.web.routers.wizard as wizard_module

    importlib.reload(wizard_module)

    from app.db import models  # noqa: F401
    from sqlmodel import SQLModel

    async with base_module.engine.begin() as conn:
        await conn.run_sync(SQLModel.metadata.create_all)

    import app.main as main_module

    importlib.reload(main_module)

    application = main_module.create_app()
    yield application, base_module, tmp_path

    await base_module.engine.dispose()


def _make_docx(paragraphs: list[tuple[str, str | None]] | None = None) -> bytes:
    """Create a minimal DOCX file in memory.

    Each entry is (text, style_name). If style_name is None, uses Normal.
    Returns raw bytes suitable for multipart upload.
    """
    doc = Document()
    if paragraphs is None:
        paragraphs = [("Test resume content", None)]
    for text, style in paragraphs:
        if style and style.startswith("Heading"):
            doc.add_heading(text, level=int(style[-1]) if style[-1].isdigit() else 1)
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


async def test_resume_upload_docx(live_app) -> None:
    app, base_module, tmp_path = live_app
    from sqlalchemy import select
    from app.db.models import Settings

    docx_bytes = _make_docx([("My resume paragraph", None)])

    async with app.router.lifespan_context(app):
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/settings/resume",
                files={"resume": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert resp.status_code == 200
            assert "uploaded successfully" in resp.text
            assert "My resume paragraph" in resp.text

            # File exists on filesystem
            resume_path = tmp_path / "resumes" / "base_resume.docx"
            assert resume_path.exists()

            # Metadata in DB
            async with base_module.async_session() as session:
                row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
                assert row.resume_filename == "resume.docx"
                assert row.resume_uploaded_at is not None


async def test_resume_upload_non_docx_rejected(live_app) -> None:
    app, _, _ = live_app
    async with app.router.lifespan_context(app):
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/settings/resume",
                files={"resume": ("notes.txt", b"plain text content", "text/plain")},
            )
            assert resp.status_code == 200
            assert "Only .docx files" in resp.text


async def test_resume_replace(live_app) -> None:
    app, _, tmp_path = live_app
    docx1 = _make_docx([("First version", None)])
    docx2 = _make_docx([("Second version", None)])

    async with app.router.lifespan_context(app):
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as client:
            await client.post(
                "/settings/resume",
                files={"resume": ("v1.docx", docx1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            resp = await client.post(
                "/settings/resume",
                files={"resume": ("v2.docx", docx2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert resp.status_code == 200
            assert "Second version" in resp.text

            # Verify only one file on disk via the resume service
            from app.resume.service import get_resume_path

            path = get_resume_path()
            assert path is not None and path.exists()
            # The file should contain the second version (replaced)
            from docx import Document as DocxDoc

            doc = DocxDoc(str(path))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            assert "Second version" in texts
            assert "First version" not in texts


async def test_resume_preview_content(live_app) -> None:
    app, _, _ = live_app
    docx_bytes = _make_docx([
        ("Experience", "Heading 1"),
        ("Senior Developer at ACME Corp", None),
        ("Education", "Heading 1"),
        ("BS Computer Science", None),
    ])
    async with app.router.lifespan_context(app):
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/settings/resume",
                files={"resume": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
            assert resp.status_code == 200
            body = resp.text
            assert "Experience" in body
            assert "Education" in body
            assert "Senior Developer at ACME Corp" in body


async def test_resume_section_no_upload(live_app) -> None:
    app, _, _ = live_app
    async with app.router.lifespan_context(app):
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.get("/settings/section/resume")
            assert resp.status_code == 200
            assert "Upload resume" in resp.text
