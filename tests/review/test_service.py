"""Phase 5 plan 05-05 Task 1 — review service unit tests.

All status writes MUST flow through ``assert_valid_transition`` so the UI
cannot push the state machine into an illegal transition. The
``approve_batch`` happy / rollback paths are the most important coverage
because a single bad row must NOT half-apply a batch approval.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.discovery.models import Job
from app.review.service import (
    approve_batch,
    approve_one,
    list_review_queue,
    reject_job,
    retailor_job,
    save_user_edits,
)
from app.tailoring.models import TailoringRecord


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _build_sample_docx(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Engineer.")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("- Built thing")
    doc.save(str(path))
    return path


async def _seed_job(
    session, *, status: str = "tailored", company: str = "Acme",
    title: str = "Engineer", score: int = 80, fingerprint_suffix: str = "1",
) -> Job:
    job = Job(
        fingerprint=f"fp-{fingerprint_suffix}",
        external_id=f"ext-{fingerprint_suffix}",
        title=title,
        company=company,
        location="",
        description="",
        url="https://example.com/j",
        source="greenhouse",
        score=score,
        status=status,
    )
    session.add(job)
    await session.commit()
    await session.refresh(job)
    return job


async def _seed_record(
    session, job: Job, *, base_path: Path, tailored_path: Path | None = None,
    version: int = 1, intensity: str = "balanced",
) -> TailoringRecord:
    rec = TailoringRecord(
        job_id=job.id,
        version=version,
        intensity=intensity,
        status="completed",
        base_resume_path=str(base_path),
        tailored_resume_path=str(tailored_path) if tailored_path else None,
    )
    session.add(rec)
    await session.commit()
    await session.refresh(rec)
    return rec


# ---------------------------------------------------------------------------
# Approve
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_approve_valid_transition_succeeds(async_session) -> None:
    job = await _seed_job(async_session, status="tailored")
    await approve_one(async_session, job.id)
    refreshed = await async_session.get(Job, job.id)
    assert refreshed.status == "approved"


@pytest.mark.asyncio
async def test_approve_from_submitted_raises_value_error(async_session) -> None:
    job = await _seed_job(async_session, status="submitted")
    with pytest.raises(ValueError):
        await approve_one(async_session, job.id)
    refreshed = await async_session.get(Job, job.id)
    assert refreshed.status == "submitted"


@pytest.mark.asyncio
async def test_batch_approve_all_or_nothing(async_session) -> None:
    j1 = await _seed_job(async_session, status="tailored", fingerprint_suffix="a")
    j2 = await _seed_job(async_session, status="tailored", fingerprint_suffix="b")
    j3 = await _seed_job(async_session, status="submitted", fingerprint_suffix="c")
    # Snapshot ids BEFORE the rollback expires/detaches the instances.
    j1_id, j2_id, j3_id = j1.id, j2.id, j3.id

    with pytest.raises(ValueError):
        await approve_batch(async_session, [j1_id, j2_id, j3_id])

    # None of them flipped — the bad row rolled back the batch.
    # `rollback()` expires every ORM-tracked attribute, so we expunge the
    # identity map and read fresh column values via Core SQL to avoid the
    # MissingGreenlet "lazy-load in sync context" trap.
    from sqlalchemy import select as _select

    async_session.expunge_all()
    rows = (
        await async_session.execute(
            _select(Job.id, Job.status).where(
                Job.id.in_([j1_id, j2_id, j3_id])
            )
        )
    ).all()
    by_id = {r.id: r.status for r in rows}
    assert by_id[j1_id] == "tailored"
    assert by_id[j2_id] == "tailored"
    assert by_id[j3_id] == "submitted"


@pytest.mark.asyncio
async def test_batch_approve_happy_path(async_session) -> None:
    j1 = await _seed_job(async_session, status="tailored", fingerprint_suffix="x")
    j2 = await _seed_job(async_session, status="tailored", fingerprint_suffix="y")
    count = await approve_batch(async_session, [j1.id, j2.id])
    assert count == 2
    for jid in (j1.id, j2.id):
        row = await async_session.get(Job, jid)
        assert row.status == "approved"


# ---------------------------------------------------------------------------
# Reject + retailor
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_reject_skip_sets_skipped(async_session) -> None:
    job = await _seed_job(async_session, status="tailored")
    await reject_job(async_session, job.id, mode="skip")
    row = await async_session.get(Job, job.id)
    assert row.status == "skipped"


@pytest.mark.asyncio
async def test_retailor_sets_status_retailoring(async_session) -> None:
    job = await _seed_job(async_session, status="tailored")
    await retailor_job(async_session, job.id)
    row = await async_session.get(Job, job.id)
    assert row.status == "retailoring"


@pytest.mark.asyncio
async def test_reject_unknown_mode_raises(async_session) -> None:
    job = await _seed_job(async_session, status="tailored")
    with pytest.raises(ValueError):
        await reject_job(async_session, job.id, mode="bogus")


# ---------------------------------------------------------------------------
# save_user_edits — manual_edit record, no LLM
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_save_user_edits_creates_manual_edit_record(
    async_session, tmp_path, env_with_fernet
) -> None:
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_v1 = _build_sample_docx(tmp_path / "v1.docx")

    job = await _seed_job(async_session, status="tailored")
    await _seed_record(
        async_session, job, base_path=base_path, tailored_path=tailored_v1, version=1
    )

    edits = {
        "sections": [
            {"heading": "Summary", "content": ["Engineer with edited bullets."]},
            {"heading": "Experience", "content": ["- Built thing now updated"]},
        ]
    }
    new_record = await save_user_edits(async_session, job.id, edits)

    assert new_record.intensity == "manual_edit"
    assert new_record.version == 2
    assert new_record.input_tokens == 0
    assert new_record.output_tokens == 0
    assert new_record.estimated_cost_dollars == 0.0
    assert new_record.error_message == "user_edit"
    assert new_record.tailored_resume_path is not None
    assert Path(new_record.tailored_resume_path).exists()
    # Job status untouched — user approves separately.
    job_row = await async_session.get(Job, job.id)
    assert job_row.status == "tailored"


# ---------------------------------------------------------------------------
# list_review_queue — sort whitelist + filter
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_list_review_queue_sorts_by_whitelisted_column(
    async_session, tmp_path
) -> None:
    base_path = _build_sample_docx(tmp_path / "base.docx")
    j1 = await _seed_job(async_session, company="Acme", fingerprint_suffix="1")
    j2 = await _seed_job(async_session, company="Zeta", fingerprint_suffix="2")
    await _seed_record(async_session, j1, base_path=base_path)
    await _seed_record(async_session, j2, base_path=base_path)

    rows, total = await list_review_queue(
        async_session, sort_by="company", sort_dir="asc"
    )
    assert total == 2
    assert [r[0].company for r in rows] == ["Acme", "Zeta"]

    rows_desc, _ = await list_review_queue(
        async_session, sort_by="company", sort_dir="desc"
    )
    assert [r[0].company for r in rows_desc] == ["Zeta", "Acme"]


@pytest.mark.asyncio
async def test_list_review_queue_rejects_unknown_sort_column(
    async_session, tmp_path
) -> None:
    base_path = _build_sample_docx(tmp_path / "base.docx")
    j1 = await _seed_job(async_session, fingerprint_suffix="1")
    await _seed_record(async_session, j1, base_path=base_path)

    # Unknown column does NOT raise — falls back to tailored_at desc.
    rows, total = await list_review_queue(async_session, sort_by="bogus_col")
    assert total == 1
    assert rows[0][0].id == j1.id


@pytest.mark.asyncio
async def test_list_review_queue_filters_by_status(
    async_session, tmp_path
) -> None:
    base_path = _build_sample_docx(tmp_path / "base.docx")
    j1 = await _seed_job(async_session, status="tailored", fingerprint_suffix="1")
    j2 = await _seed_job(async_session, status="submitted", fingerprint_suffix="2")
    await _seed_record(async_session, j1, base_path=base_path)
    await _seed_record(async_session, j2, base_path=base_path)

    rows, total = await list_review_queue(async_session)  # default filter
    ids = {r[0].id for r in rows}
    assert j1.id in ids
    assert j2.id not in ids
    assert total == 1
