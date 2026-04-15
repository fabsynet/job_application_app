"""Phase 5 plan 05-05 Task 2 — review router integration tests.

Mirrors the live_app fixture pattern used by tests/integration/* — reload
``app.config`` + ``app.db.base`` against a fresh tmp DATA_DIR so each test
case runs against an isolated SQLite file. We then build the FastAPI app
inside the lifespan context and drive it through httpx.AsyncClient.
"""
from __future__ import annotations

import importlib
from pathlib import Path

import httpx
import pytest
import pytest_asyncio
from cryptography.fernet import Fernet
from docx import Document
from sqlalchemy import select


@pytest_asyncio.fixture
async def live_app(monkeypatch: pytest.MonkeyPatch, tmp_path: Path):
    key = Fernet.generate_key().decode()
    monkeypatch.setenv("FERNET_KEY", key)
    monkeypatch.setenv("TZ", "UTC")
    monkeypatch.setenv("BIND_ADDRESS", "127.0.0.1")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("LOG_LEVEL", "INFO")

    import app.config as config_module

    config_module.get_settings.cache_clear()
    importlib.reload(config_module)

    import app.db.base as base_module

    importlib.reload(base_module)

    from app.db import models  # noqa: F401
    from sqlmodel import SQLModel

    async with base_module.engine.begin() as conn:
        await conn.run_sync(SQLModel.metadata.create_all)

    import app.main as main_module

    importlib.reload(main_module)

    app_obj = main_module.create_app()
    yield app_obj, base_module, key, tmp_path

    await base_module.engine.dispose()


def _build_sample_docx(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Engineer.")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("- Built thing")
    doc.save(str(path))
    return path


async def _seed_job_with_record(
    base_module,
    *,
    company: str,
    title: str,
    status: str,
    fingerprint: str,
    base_path: Path,
    tailored_path: Path | None = None,
    intensity: str = "balanced",
):
    from app.discovery.models import Job
    from app.tailoring.models import TailoringRecord

    async with base_module.async_session() as session:
        job = Job(
            fingerprint=fingerprint,
            external_id=fingerprint,
            title=title,
            company=company,
            location="",
            description="",
            url="https://example.com/j",
            source="greenhouse",
            score=85,
            status=status,
        )
        session.add(job)
        await session.commit()
        await session.refresh(job)

        rec = TailoringRecord(
            job_id=job.id,
            version=1,
            intensity=intensity,
            status="completed",
            base_resume_path=str(base_path),
            tailored_resume_path=str(tailored_path) if tailored_path else None,
        )
        session.add(rec)
        await session.commit()
        await session.refresh(rec)
        return job.id, rec.id


async def test_review_index_renders(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx")
    await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Senior Engineer",
        status="tailored",
        fingerprint="fp-idx-1",
        base_path=base_path,
        tailored_path=tailored_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get("/review")
            assert resp.status_code == 200
            assert "<table" in resp.text
            assert "Acme" in resp.text
            assert "Senior Engineer" in resp.text


async def test_review_drawer_returns_diff(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Stripe",
        title="Backend Engineer",
        status="tailored",
        fingerprint="fp-draw-1",
        base_path=base_path,
        tailored_path=tailored_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get(
                f"/review/{job_id}",
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200
            # Drawer markup
            assert "Stripe" in resp.text
            assert "Edit tailored sections" in resp.text
            # Edit textareas
            assert "section_" in resp.text
            assert "heading_" in resp.text


async def test_review_save_edits_creates_manual_edit_record(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Engineer",
        status="tailored",
        fingerprint="fp-save-1",
        base_path=base_path,
        tailored_path=tailored_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                f"/review/{job_id}/save-edits",
                data={
                    "heading_0": "Summary",
                    "section_0": "Edited bullet line one\nEdited bullet line two",
                    "heading_1": "Experience",
                    "section_1": "- Did new thing",
                },
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200
            assert "Saved" in resp.text

    # Verify a new manual_edit record exists.
    from app.tailoring.models import TailoringRecord

    async with base_module.async_session() as session:
        rows = (
            await session.execute(
                select(TailoringRecord).where(TailoringRecord.job_id == job_id)
            )
        ).scalars().all()
        assert len(rows) == 2
        latest = max(rows, key=lambda r: r.version)
        assert latest.intensity == "manual_edit"
        assert latest.input_tokens == 0


async def test_review_approve_one_flips_status(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="tailored",
        fingerprint="fp-app-1",
        base_path=base_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                f"/review/{job_id}/approve",
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200

    from app.discovery.models import Job

    async with base_module.async_session() as session:
        row = (
            await session.execute(select(Job).where(Job.id == job_id))
        ).scalar_one()
        assert row.status == "approved"


async def test_review_approve_illegal_returns_422_toast(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="submitted",  # cannot be approved
        fingerprint="fp-bad-1",
        base_path=base_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                f"/review/{job_id}/approve",
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 422
            assert "toast" in resp.text.lower()

    from app.discovery.models import Job

    async with base_module.async_session() as session:
        row = (
            await session.execute(select(Job).where(Job.id == job_id))
        ).scalar_one()
        assert row.status == "submitted"  # unchanged


async def test_review_confirm_batch_modal(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    j1, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng A",
        status="tailored",
        fingerprint="fp-cb-1",
        base_path=base_path,
    )
    j2, _ = await _seed_job_with_record(
        base_module,
        company="Zeta",
        title="Eng B",
        status="tailored",
        fingerprint="fp-cb-2",
        base_path=base_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get(
                f"/review/confirm-approve?job_ids={j1}&job_ids={j2}",
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200
            assert "Acme" in resp.text
            assert "Zeta" in resp.text
            assert "Approve 2" in resp.text


async def test_review_approve_batch(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    j1, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="A",
        status="tailored",
        fingerprint="fp-bat-1",
        base_path=base_path,
    )
    j2, _ = await _seed_job_with_record(
        base_module,
        company="Beta",
        title="B",
        status="tailored",
        fingerprint="fp-bat-2",
        base_path=base_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                "/review/approve-batch",
                data={"job_ids": [str(j1), str(j2)]},
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200

    from app.discovery.models import Job

    async with base_module.async_session() as session:
        for jid in (j1, j2):
            row = (
                await session.execute(select(Job).where(Job.id == jid))
            ).scalar_one()
            assert row.status == "approved"


async def test_review_reject_skip(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="tailored",
        fingerprint="fp-rj-1",
        base_path=base_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                f"/review/{job_id}/reject",
                data={"mode": "skip"},
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200

    from app.discovery.models import Job

    async with base_module.async_session() as session:
        row = (
            await session.execute(select(Job).where(Job.id == job_id))
        ).scalar_one()
        assert row.status == "skipped"


async def test_review_reject_retailor(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="tailored",
        fingerprint="fp-rj-2",
        base_path=base_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                f"/review/{job_id}/reject",
                data={"mode": "retailor"},
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200

    from app.discovery.models import Job

    async with base_module.async_session() as session:
        row = (
            await session.execute(select(Job).where(Job.id == job_id))
        ).scalar_one()
        assert row.status == "retailoring"


async def test_review_drawer_after_save_shows_edits(live_app) -> None:
    """After save_user_edits, re-opening the drawer reads the NEW record's
    DOCX so the textareas reflect the user's edits."""
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="tailored",
        fingerprint="fp-after-1",
        base_path=base_path,
        tailored_path=tailored_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            await client.post(
                f"/review/{job_id}/save-edits",
                data={
                    "heading_0": "Summary",
                    "section_0": "Brand new summary line",
                    "heading_1": "Experience",
                    "section_1": "- Brand new bullet",
                },
                headers={"HX-Request": "true"},
            )
            resp = await client.get(
                f"/review/{job_id}",
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200
            assert "Brand new" in resp.text
