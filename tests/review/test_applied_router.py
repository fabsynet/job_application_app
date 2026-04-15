"""Phase 5 plan 05-08 Task 2 — applied router integration tests.

Mirrors the live_app fixture pattern used by ``tests/review/test_router.py``
(reload ``app.config`` + ``app.db.base`` against a fresh tmp DATA_DIR per
test so SQLite isolation is clean).
"""
from __future__ import annotations

import importlib
from pathlib import Path

import httpx
import pytest
import pytest_asyncio
from cryptography.fernet import Fernet
from docx import Document
from sqlalchemy import select


@pytest_asyncio.fixture
async def live_app(monkeypatch: pytest.MonkeyPatch, tmp_path: Path):
    key = Fernet.generate_key().decode()
    monkeypatch.setenv("FERNET_KEY", key)
    monkeypatch.setenv("TZ", "UTC")
    monkeypatch.setenv("BIND_ADDRESS", "127.0.0.1")
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    monkeypatch.setenv("LOG_LEVEL", "INFO")

    import app.config as config_module

    config_module.get_settings.cache_clear()
    importlib.reload(config_module)

    import app.db.base as base_module

    importlib.reload(base_module)

    from app.db import models  # noqa: F401
    from sqlmodel import SQLModel

    async with base_module.engine.begin() as conn:
        await conn.run_sync(SQLModel.metadata.create_all)

    import app.main as main_module

    importlib.reload(main_module)

    app_obj = main_module.create_app()
    yield app_obj, base_module, key, tmp_path

    await base_module.engine.dispose()


def _build_sample_docx(path: Path, body: str = "Sample body") -> Path:
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(body)
    doc.save(str(path))
    return path


async def _seed_job_with_record(
    base_module,
    *,
    company: str,
    title: str,
    status: str,
    fingerprint: str,
    base_path: Path,
    tailored_path: Path | None = None,
    cover_letter_path: Path | None = None,
    source: str = "greenhouse",
):
    from app.discovery.models import Job
    from app.tailoring.models import TailoringRecord

    async with base_module.async_session() as session:
        job = Job(
            fingerprint=fingerprint,
            external_id=fingerprint,
            title=title,
            company=company,
            location="",
            description="A role about Python and Postgres.",
            description_html="<p>A role.</p>",
            url="https://example.com/j",
            source=source,
            score=85,
            status=status,
        )
        session.add(job)
        await session.commit()
        await session.refresh(job)

        rec = TailoringRecord(
            job_id=job.id,
            version=1,
            intensity="balanced",
            status="completed",
            base_resume_path=str(base_path),
            tailored_resume_path=str(tailored_path) if tailored_path else None,
            cover_letter_path=str(cover_letter_path) if cover_letter_path else None,
        )
        session.add(rec)
        await session.commit()
        await session.refresh(rec)
        return job.id, rec.id


async def _seed_submission(base_module, *, job_id: int, tailoring_record_id: int):
    from datetime import datetime

    from app.submission.models import Submission

    async with base_module.async_session() as session:
        sub = Submission(
            job_id=job_id,
            tailoring_record_id=tailoring_record_id,
            status="sent",
            smtp_from="me@example.com",
            smtp_to="jobs@example.com",
            subject="Application for role",
            attachment_filename="me_Acme_Resume.docx",
            sent_at=datetime.utcnow(),
        )
        session.add(sub)
        await session.commit()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


async def test_get_applied_index_renders_counts(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx", body="Tailored.")
    j1, r1 = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Senior Engineer",
        status="submitted",
        fingerprint="fp-s1",
        base_path=base_path,
        tailored_path=tailored_path,
    )
    j2, r2 = await _seed_job_with_record(
        base_module,
        company="Beta",
        title="Backend Engineer",
        status="submitted",
        fingerprint="fp-s2",
        base_path=base_path,
        tailored_path=tailored_path,
    )
    j3, _ = await _seed_job_with_record(
        base_module,
        company="Gamma",
        title="SRE",
        status="failed",
        fingerprint="fp-f1",
        base_path=base_path,
        tailored_path=tailored_path,
    )
    await _seed_submission(base_module, job_id=j1, tailoring_record_id=r1)
    await _seed_submission(base_module, job_id=j2, tailoring_record_id=r2)

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get("/applied")
            assert resp.status_code == 200
            assert "Applied Jobs" in resp.text
            assert "Acme" in resp.text
            assert "Beta" in resp.text
            assert "Gamma" in resp.text
            assert "submitted" in resp.text
            # Counts badges rendered
            assert "Today" in resp.text
            assert "Last 7 days" in resp.text


async def test_get_applied_detail_renders_all_fields(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx", body="Tailored content xyz.")
    cover_path = _build_sample_docx(tmp_path / "cover.docx", body="Dear hiring manager,\n\nRegards.")
    job_id, rec_id = await _seed_job_with_record(
        base_module,
        company="Stripe",
        title="Backend Engineer",
        status="submitted",
        fingerprint="fp-d1",
        base_path=base_path,
        tailored_path=tailored_path,
        cover_letter_path=cover_path,
    )
    await _seed_submission(base_module, job_id=job_id, tailoring_record_id=rec_id)

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get(f"/applied/{job_id}")
            assert resp.status_code == 200
            assert "Stripe" in resp.text
            assert "Backend Engineer" in resp.text
            # Submission metadata
            assert "jobs@example.com" in resp.text
            assert "me_Acme_Resume.docx" in resp.text
            # Cover letter text
            assert "hiring manager" in resp.text.lower()
            # Tailored preview was embedded
            assert "Tailored content" in resp.text


async def test_download_resume_for_submitted_job(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx")
    job_id, rec_id = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="submitted",
        fingerprint="fp-dls",
        base_path=base_path,
        tailored_path=tailored_path,
    )
    await _seed_submission(base_module, job_id=job_id, tailoring_record_id=rec_id)

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get(f"/applied/{job_id}/download")
            assert resp.status_code == 200
            assert "wordprocessingml" in resp.headers["content-type"]


async def test_download_resume_for_approved_unsent_job(live_app) -> None:
    """SC-5 explicit requirement: approved-but-unsent jobs must download."""
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    tailored_path = _build_sample_docx(tmp_path / "v1.docx")
    job_id, _ = await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng",
        status="approved",  # NOT submitted
        fingerprint="fp-dla",
        base_path=base_path,
        tailored_path=tailored_path,
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get(f"/applied/{job_id}/download")
            assert resp.status_code == 200
            assert "wordprocessingml" in resp.headers["content-type"]


async def test_banner_shows_when_rate_limited(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    from app.db.models import Run

    async with base_module.async_session() as session:
        run = Run(
            status="succeeded",
            counts={"rate_limited": True, "submitted": 3, "approved": 2},
        )
        session.add(run)
        await session.commit()

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get("/applied")
            assert resp.status_code == 200
            assert "Daily cap hit" in resp.text
            assert "Raise cap" in resp.text


async def test_raise_cap_increments_setting(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    from app.db.models import Settings

    # Seed a starting cap
    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one_or_none()
        if row is None:
            row = Settings(id=1, daily_cap=20)
            session.add(row)
        else:
            row.daily_cap = 20
        await session.commit()

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                "/applied/raise-cap",
                data={"raise_by": 5},
                headers={"HX-Request": "true"},
            )
            assert resp.status_code == 200

    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
        assert row.daily_cap == 25


async def test_settings_post_notification_email_persists(live_app) -> None:
    app, base_module, _, _ = live_app
    from app.db.models import Settings

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                "/settings/notification-email",
                data={"notification_email": "alerts@example.com"},
            )
            assert resp.status_code == 200

    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
        assert row.notification_email == "alerts@example.com"


async def test_settings_post_base_url_persists(live_app) -> None:
    app, base_module, _, _ = live_app
    from app.db.models import Settings

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                "/settings/base-url",
                data={"base_url": "https://apply.example.com"},
            )
            assert resp.status_code == 200

    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
        assert row.base_url == "https://apply.example.com"


async def test_settings_post_submissions_paused_persists(live_app) -> None:
    app, base_module, _, _ = live_app
    from app.db.models import Settings

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.post(
                "/settings/submissions-paused",
                data={"submissions_paused": "true"},
            )
            assert resp.status_code == 200

    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
        assert row.submissions_paused is True


async def test_settings_post_auto_holdout_margin_clamps(live_app) -> None:
    app, base_module, _, _ = live_app
    from app.db.models import Settings

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            # Above 50 → clamped to 50.
            resp = await client.post(
                "/settings/auto-holdout-margin",
                data={"auto_holdout_margin_pct": 999},
            )
            assert resp.status_code == 200

    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
        assert row.auto_holdout_margin_pct == 50

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            # Negative → clamped to 0.
            resp = await client.post(
                "/settings/auto-holdout-margin",
                data={"auto_holdout_margin_pct": -5},
            )
            assert resp.status_code == 200

    async with base_module.async_session() as session:
        row = (await session.execute(select(Settings).where(Settings.id == 1))).scalar_one()
        assert row.auto_holdout_margin_pct == 0


async def test_applied_list_filters_by_source_and_status(live_app) -> None:
    app, base_module, _, tmp_path = live_app
    base_path = _build_sample_docx(tmp_path / "base.docx")
    await _seed_job_with_record(
        base_module,
        company="Acme",
        title="Eng A",
        status="submitted",
        fingerprint="fp-f-g1",
        base_path=base_path,
        source="greenhouse",
    )
    await _seed_job_with_record(
        base_module,
        company="Beta",
        title="Eng B",
        status="submitted",
        fingerprint="fp-f-l1",
        base_path=base_path,
        source="lever",
    )

    async with app.router.lifespan_context(app):
        transport = httpx.ASGITransport(app=app)
        async with httpx.AsyncClient(
            transport=transport, base_url="http://test"
        ) as client:
            resp = await client.get(
                "/applied",
                params={"source": "greenhouse"},
            )
            assert resp.status_code == 200
            assert "Acme" in resp.text
            assert "Beta" not in resp.text
