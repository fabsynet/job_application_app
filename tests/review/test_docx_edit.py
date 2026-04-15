"""Phase 5 plan 05-05 Task 1 — round-trip a tailored DOCX through user edits."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.review.docx_edit import apply_user_edits, extract_sections_from_docx


def _build_sample_docx(path: Path) -> Path:
    """Create a small fixture DOCX with 3 headings and bullets each."""
    doc = Document()
    doc.add_paragraph("Jane Doe")  # pre-heading name line
    doc.add_paragraph("jane@example.com")  # pre-heading contact line

    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Senior engineer with FastAPI experience.")

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("- Built feature X")
    doc.add_paragraph("- Built feature Y")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("BS Computer Science")

    doc.save(str(path))
    return path


def test_extract_sections_round_trips(tmp_path: Path) -> None:
    docx_path = _build_sample_docx(tmp_path / "tailored.docx")
    extracted = extract_sections_from_docx(docx_path)
    sections = extracted["sections"]

    # Pre-heading bucket + 3 real sections.
    headings = [s["heading"] for s in sections]
    assert "Summary" in headings
    assert "Experience" in headings
    assert "Education" in headings

    by_heading = {s["heading"]: s for s in sections}
    assert by_heading["Summary"]["content"] == [
        "Senior engineer with FastAPI experience."
    ]
    assert by_heading["Experience"]["content"] == [
        "- Built feature X",
        "- Built feature Y",
    ]
    assert by_heading["Education"]["content"] == ["BS Computer Science"]


def test_extract_preserves_empty_preamble_section(tmp_path: Path) -> None:
    docx_path = _build_sample_docx(tmp_path / "tailored2.docx")
    extracted = extract_sections_from_docx(docx_path)
    first = extracted["sections"][0]
    # Pre-heading paragraphs (name + email) belong to a leading section
    # whose heading is the empty string.
    assert first["heading"] == ""
    assert "Jane Doe" in first["content"]
    assert "jane@example.com" in first["content"]


def test_apply_user_edits_uses_build_tailored_docx(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """`apply_user_edits` MUST go through `build_tailored_docx` (the only
    safe DOCX mutator)."""
    base_path = _build_sample_docx(tmp_path / "base.docx")
    out_path = tmp_path / "out.docx"

    captured: dict = {}

    def _spy(*, base_resume_path, tailored_sections, output_path):
        captured["base_resume_path"] = base_resume_path
        captured["tailored_sections"] = tailored_sections
        captured["output_path"] = output_path
        # Touch the output file so callers see a file at the path.
        Path(output_path).write_bytes(b"stub")
        return Path(output_path)

    import app.tailoring.docx_writer as dw

    monkeypatch.setattr(dw, "build_tailored_docx", _spy)

    edits = {"sections": [{"heading": "Summary", "content": ["edited line"]}]}
    result = apply_user_edits(
        base_resume_path=base_path,
        edited_sections=edits,
        output_path=out_path,
    )
    assert result == out_path
    assert captured["base_resume_path"] == base_path
    assert captured["output_path"] == out_path
    assert captured["tailored_sections"] == edits
