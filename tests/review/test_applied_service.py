"""Phase 5 plan 05-08 Task 1 — applied service unit tests.

Covers counts by window, the filterable/sortable list (including
approved-but-unsent jobs per SC-5), the detail payload, and the
``applied_artifact_paths`` helper.
"""
from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path

import pytest
from docx import Document

from app.discovery.models import Job
from app.review.applied_service import (
    APPLIED_SORT_COLUMNS,
    applied_artifact_paths,
    get_applied_detail,
    list_applied_jobs,
    state_counts_for_window,
)
from app.submission.models import Submission
from app.tailoring.models import TailoringRecord


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _build_sample_docx(path: Path, *, body: str = "Engineer.") -> Path:
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(body)
    doc.save(str(path))
    return path


async def _seed_job(
    session,
    *,
    status: str,
    company: str = "Acme",
    title: str = "Engineer",
    source: str = "greenhouse",
    fingerprint_suffix: str = "1",
    first_seen_at: datetime | None = None,
    score: int = 80,
) -> Job:
    job = Job(
        fingerprint=f"fp-{fingerprint_suffix}",
        external_id=f"ext-{fingerprint_suffix}",
        title=title,
        company=company,
        location="",
        description="A role.",
        url="https://example.com/j",
        source=source,
        score=score,
        status=status,
        first_seen_at=first_seen_at or datetime.utcnow(),
    )
    session.add(job)
    await session.commit()
    await session.refresh(job)
    return job


async def _seed_submission(
    session,
    *,
    job_id: int,
    tailoring_record_id: int,
    sent_at: datetime | None = None,
    status: str = "sent",
) -> Submission:
    sub = Submission(
        job_id=job_id,
        tailoring_record_id=tailoring_record_id,
        status=status,
        smtp_from="me@example.com",
        smtp_to="jobs@example.com",
        subject="Application for role",
        attachment_filename="me_Acme_Resume.docx",
        sent_at=sent_at,
    )
    session.add(sub)
    await session.commit()
    await session.refresh(sub)
    return sub


async def _seed_record(
    session,
    *,
    job_id: int,
    base_path: Path,
    tailored_path: Path | None = None,
    cover_letter_path: Path | None = None,
    version: int = 1,
) -> TailoringRecord:
    rec = TailoringRecord(
        job_id=job_id,
        version=version,
        intensity="balanced",
        status="completed",
        base_resume_path=str(base_path),
        tailored_resume_path=str(tailored_path) if tailored_path else None,
        cover_letter_path=str(cover_letter_path) if cover_letter_path else None,
    )
    session.add(rec)
    await session.commit()
    await session.refresh(rec)
    return rec


# ---------------------------------------------------------------------------
# state_counts_for_window
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_state_counts_groups_by_status(async_session) -> None:
    await _seed_job(async_session, status="submitted", fingerprint_suffix="a")
    await _seed_job(async_session, status="submitted", fingerprint_suffix="b")
    await _seed_job(async_session, status="failed", fingerprint_suffix="c")
    await _seed_job(async_session, status="approved", fingerprint_suffix="d")

    since = datetime.utcnow() - timedelta(hours=1)
    counts = await state_counts_for_window(async_session, since=since)
    assert counts.submitted == 2
    assert counts.failed == 1
    assert counts.approved == 1
    assert counts.skipped == 0


@pytest.mark.asyncio
async def test_state_counts_respects_window(async_session) -> None:
    # One seeded "yesterday", one seeded now.
    yesterday = datetime.utcnow() - timedelta(days=1, hours=2)
    await _seed_job(
        async_session,
        status="submitted",
        fingerprint_suffix="old",
        first_seen_at=yesterday,
    )
    await _seed_job(async_session, status="submitted", fingerprint_suffix="new")

    since = datetime.utcnow() - timedelta(hours=1)
    counts = await state_counts_for_window(async_session, since=since)
    assert counts.submitted == 1  # only the fresh row


# ---------------------------------------------------------------------------
# list_applied_jobs
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_list_applied_includes_approved_but_unsent(async_session, tmp_path) -> None:
    """SC-5: approved-but-unsent jobs must appear in the applied table."""
    base = _build_sample_docx(tmp_path / "base.docx")
    job = await _seed_job(async_session, status="approved", fingerprint_suffix="ap1")
    await _seed_record(async_session, job_id=job.id, base_path=base)

    rows, total = await list_applied_jobs(async_session)
    assert total == 1
    assert len(rows) == 1
    j, sub, rec = rows[0]
    assert j.id == job.id
    assert sub is None  # no submission row yet
    assert rec is not None  # tailoring record present


@pytest.mark.asyncio
async def test_list_applied_sorted_by_submitted_at_desc(async_session, tmp_path) -> None:
    base = _build_sample_docx(tmp_path / "base.docx")
    # Three submitted jobs with sent_at offsets.
    now = datetime.utcnow()
    ids: list[int] = []
    for i in range(3):
        job = await _seed_job(
            async_session,
            status="submitted",
            fingerprint_suffix=f"s{i}",
            company=f"Co{i}",
        )
        rec = await _seed_record(async_session, job_id=job.id, base_path=base)
        await _seed_submission(
            async_session,
            job_id=job.id,
            tailoring_record_id=rec.id,
            sent_at=now - timedelta(hours=i),
        )
        ids.append(job.id)

    rows, _ = await list_applied_jobs(async_session)
    # Most-recent sent_at first — that's ids[0].
    sent_at_values = [r[1].sent_at for r in rows if r[1] is not None]
    assert sent_at_values == sorted(sent_at_values, reverse=True)
    assert rows[0][0].id == ids[0]


@pytest.mark.asyncio
async def test_list_applied_filtered_by_source(async_session, tmp_path) -> None:
    await _seed_job(async_session, status="submitted", source="greenhouse", fingerprint_suffix="g1")
    await _seed_job(async_session, status="submitted", source="greenhouse", fingerprint_suffix="g2")
    await _seed_job(async_session, status="submitted", source="lever", fingerprint_suffix="l1")

    rows, total = await list_applied_jobs(
        async_session, source_filter=["greenhouse"]
    )
    assert total == 2
    assert all(r[0].source == "greenhouse" for r in rows)


@pytest.mark.asyncio
async def test_list_applied_unknown_sort_column_falls_back(async_session) -> None:
    await _seed_job(async_session, status="submitted", fingerprint_suffix="x")
    # ``random`` is not in APPLIED_SORT_COLUMNS — should not raise.
    rows, total = await list_applied_jobs(async_session, sort_by="random")
    assert total == 1
    assert len(rows) == 1


# ---------------------------------------------------------------------------
# get_applied_detail
# ---------------------------------------------------------------------------


@pytest.mark.asyncio
async def test_get_applied_detail_returns_paths_and_cover_letter(
    async_session, tmp_path
) -> None:
    base = _build_sample_docx(tmp_path / "base.docx")
    tailored = _build_sample_docx(tmp_path / "v1.docx", body="Tailored content.")
    cover = _build_sample_docx(
        tmp_path / "cover.docx", body="Dear hiring manager,\n\nThank you."
    )
    job = await _seed_job(async_session, status="submitted", fingerprint_suffix="d1")
    rec = await _seed_record(
        async_session,
        job_id=job.id,
        base_path=base,
        tailored_path=tailored,
        cover_letter_path=cover,
    )
    await _seed_submission(
        async_session,
        job_id=job.id,
        tailoring_record_id=rec.id,
        sent_at=datetime.utcnow(),
    )

    data = await get_applied_detail(async_session, job.id)
    assert data["job"].id == job.id
    assert data["record"].id == rec.id
    assert data["submission"] is not None
    assert "hiring manager" in data["cover_letter_text"].lower()
    # Preview should contain the tailored body we seeded.
    assert "Tailored content" in data["tailored_preview_html"]


@pytest.mark.asyncio
async def test_get_applied_detail_missing_job_returns_empty(async_session) -> None:
    data = await get_applied_detail(async_session, 9999)
    assert data == {}


# ---------------------------------------------------------------------------
# applied_artifact_paths
# ---------------------------------------------------------------------------


def test_applied_artifact_paths_none_safe() -> None:
    out = applied_artifact_paths(None)
    assert out == {"resume": None, "cover_letter": None}


def test_applied_artifact_paths_populated(tmp_path) -> None:
    resume = tmp_path / "r.docx"
    cover = tmp_path / "c.docx"
    rec = TailoringRecord(
        job_id=1,
        version=1,
        intensity="balanced",
        status="completed",
        base_resume_path=str(tmp_path / "base.docx"),
        tailored_resume_path=str(resume),
        cover_letter_path=str(cover),
    )
    out = applied_artifact_paths(rec)
    assert out["resume"] == Path(resume)
    assert out["cover_letter"] == Path(cover)


def test_applied_sort_columns_has_submitted_at() -> None:
    assert "submitted_at" in APPLIED_SORT_COLUMNS
    assert "company" in APPLIED_SORT_COLUMNS
